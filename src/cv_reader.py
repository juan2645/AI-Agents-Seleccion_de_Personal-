import os
import glob
from typing import List, Dict, Any
from docx import Document
import logging

class CVReaderAgent:
    """Clase para leer CVs desde diferentes formatos de archivo"""
    
    def __init__(self, cv_folder: str = "curriculums"):
        self.cv_folder = cv_folder
        self.supported_extensions = ['.docx', '.doc', '.txt', '.pdf']
        
    def read_word_document(self, file_path: str) -> str:
        """Lee un archivo Word (.docx) y extrae el texto"""
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text.strip())
            return '\n'.join(text)
        except Exception as e:
            logging.error(f"Error leyendo archivo Word {file_path}: {str(e)}")
            return ""
    
    def read_text_file(self, file_path: str) -> str:
        """Lee un archivo de texto"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logging.error(f"Error leyendo archivo de texto {file_path}: {str(e)}")
            return ""
    
    def get_cv_files(self) -> List[str]:
        """Obtiene la lista de archivos de CV en la carpeta"""
        cv_files = []
        
        if not os.path.exists(self.cv_folder):
            logging.warning(f"La carpeta {self.cv_folder} no existe")
            return cv_files
        
        # Buscar archivos con extensiones soportadas
        for ext in self.supported_extensions:
            pattern = os.path.join(self.cv_folder, f"*{ext}")
            cv_files.extend(glob.glob(pattern))
        
        return sorted(cv_files)
    
    def read_cv_file(self, file_path: str) -> Dict[str, Any]:
        """Lee un archivo de CV y retorna información del candidato"""
        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Extraer texto según el tipo de archivo
        if file_ext == '.docx':
            text = self.read_word_document(file_path)
        elif file_ext in ['.txt', '.doc']:
            text = self.read_text_file(file_path)
        else:
            logging.warning(f"Formato de archivo no soportado: {file_ext}")
            text = ""
        
        return {
            'file_name': file_name,
            'file_path': file_path,
            'text': text,
            'file_size': os.path.getsize(file_path) if os.path.exists(file_path) else 0
        }
    
    def read_all_cvs(self) -> List[Dict[str, Any]]:
        """Lee todos los CVs en la carpeta"""
        cv_files = self.get_cv_files()
        cvs = []
        
        print(f"📁 Buscando CVs en la carpeta: {self.cv_folder}")
        print(f"📄 Archivos encontrados: {len(cv_files)}")
        
        for file_path in cv_files:
            print(f"  📖 Leyendo: {os.path.basename(file_path)}")
            cv_data = self.read_cv_file(file_path)
            if cv_data['text']:
                cvs.append(cv_data)
                print(f"    ✅ Extraído: {len(cv_data['text'])} caracteres")
            else:
                print(f"    ❌ Error al leer el archivo")
        
        return cvs
    
    def get_cv_texts(self) -> List[str]:
        """Obtiene solo los textos de los CVs"""
        cvs = self.read_all_cvs()
        return [cv['text'] for cv in cvs if cv['text']]
    
    def create_sample_cv(self, filename: str, content: str):
        """Crea un CV de ejemplo en la carpeta"""
        if not os.path.exists(self.cv_folder):
            os.makedirs(self.cv_folder)
        
        file_path = os.path.join(self.cv_folder, filename)
        
        if filename.endswith('.docx'):
            # Crear documento Word
            doc = Document()
            paragraphs = content.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            doc.save(file_path)
        else:
            # Crear archivo de texto
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(content)
        
        print(f"✅ CV de ejemplo creado: {file_path}")
